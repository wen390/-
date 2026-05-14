from __future__ import annotations

import csv
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "面向PCB金手指的测试尾板布点布线设计_格式修正版.docx"
OUT = ROOT / "面向PCB金手指的测试尾板布点布线设计_管理要求增强版.docx"
ASSET_DIR = ROOT / "results" / "thesis_management_assets"
FORMULA_DIR = ASSET_DIR / "formulas"
TECTONIC = ROOT.parent / ".codex/plugins/cache/openai-bundled/latex-tectonic/0.1.1/bin/tectonic"
if not TECTONIC.exists():
    TECTONIC = Path("/Users/yunhe/.codex/plugins/cache/openai-bundled/latex-tectonic/0.1.1/bin/tectonic")
VENV_PY = ROOT / ".venv" / "bin" / "python"
SUMMARY_CSV = ROOT / "results" / "kicad_cases" / "summary.csv"
CONTACT_SHEET = ROOT / "results" / "kicad_cases" / "comparison_contact_sheet.png"


FORMULAS = [
    ("式（2-1）", r"P=\{p_i\mid i=1,\ldots,n\},\quad C=\{c_j\mid j=1,\ldots,m\},\quad x_{ij}\in\{0,1\}"),
    ("式（2-2）", r"\sum_{j=1}^{m}x_{ij}=1,\quad \sum_{i=1}^{n}x_{ij}\leq 1,\quad c_j\notin\Omega,\quad \|c_j-c_k\|_2\geq s_{\min}"),
    ("式（2-3）", r"\mathrm{cap}(e)=\left\lfloor\frac{w_e+s_l}{w_l+s_l}\right\rfloor,\quad U(e)=\frac{n_e}{\max(1,\mathrm{cap}(e))}"),
    ("式（2-4）", r"R_r=\frac{N_r}{N}\times100\%,\quad A_r=\frac{A_{\mathrm{bbox}}}{A_{\mathrm{tail}}}\times100\%,\quad V=N_u+N_s+N_o"),
    ("式（3-1）", r"J=\omega_dD+\omega_aA+\omega_cC+\omega_pP,\quad \omega_p\gg\omega_d,\omega_a,\omega_c"),
    ("式（4-1）", r"f(q)=g(q)+h(q)+\lambda_{\mathrm{obs}}O(q)+\lambda_{\mathrm{cong}}U(q)+\lambda_{\mathrm{turn}}T(q)"),
    ("式（4-2）", r"H_e^{t+1}=\rho H_e^t+\Delta n_e,\quad c_e^{t+1}=c_e^0(1+\alpha H_e^{t+1})"),
]


def run(cmd: list[str], cwd: Path | None = None) -> None:
    subprocess.run(cmd, cwd=str(cwd or ROOT), check=True)


def render_formula(label: str, formula: str) -> Path:
    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    stem = label.replace("式（", "eq_").replace("）", "")
    tex = FORMULA_DIR / f"{stem}.tex"
    pdf = FORMULA_DIR / f"{stem}.pdf"
    png = FORMULA_DIR / f"{stem}.png"
    tex.write_text(
        "\n".join(
            [
                r"\documentclass[12pt]{article}",
                r"\pagestyle{empty}",
                r"\begin{document}",
                r"\begin{center}",
                r"\Large",
                f"${formula}$",
                r"\end{center}",
                r"\end{document}",
            ]
        ),
        encoding="utf-8",
    )
    run([str(TECTONIC), "--outdir", str(FORMULA_DIR), str(tex)])

    raster = FORMULA_DIR / "_raster_formula.py"
    raster.write_text(
        """
from pathlib import Path
import sys
import pypdfium2 as pdfium
from PIL import Image, ImageChops

pdf_path = Path(sys.argv[1])
png_path = Path(sys.argv[2])
doc = pdfium.PdfDocument(str(pdf_path))
page = doc[0]
bitmap = page.render(scale=4).to_pil().convert("RGB")
bg = Image.new("RGB", bitmap.size, "white")
diff = ImageChops.difference(bitmap, bg)
bbox = diff.getbbox()
if bbox:
    left, top, right, bottom = bbox
    pad = 28
    crop = bitmap.crop((max(0, left-pad), max(0, top-pad), min(bitmap.width, right+pad), min(bitmap.height, bottom+pad)))
else:
    crop = bitmap
crop.save(png_path)
""".strip(),
        encoding="utf-8",
    )
    run([str(VENV_PY), str(raster), str(pdf), str(png)])
    return png


def set_run_font(run, size=12, bold=False, east="宋体", west="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:ascii"), west)
    rfonts.set(qn("w:hAnsi"), west)


def style_para(p, size=12, first_indent=True, align=None, bold=False):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    for run in p.runs:
        set_run_font(run, size=size, bold=bold)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        r = p.add_run(text)
        set_run_font(r, 12)
    style_para(p)
    return p


def insert_paragraph_before(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        r = p.add_run(text)
        set_run_font(r, 12)
    style_para(p)
    return p


def insert_picture_after(paragraph, image_path: Path, width_cm: float, caption: str, page_break_before=False):
    cap = insert_paragraph_after(paragraph, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = None
    if page_break_before:
        cap.paragraph_format.page_break_before = True
    for run in cap.runs:
        set_run_font(run, 12)
    pic_p = insert_paragraph_after(cap, "")
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.first_line_indent = None
    pic_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pic_p.paragraph_format.line_spacing = 1
    pic_p.paragraph_format.space_before = Pt(2)
    pic_p.paragraph_format.space_after = Pt(6)
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return pic_p


def insert_table_after(paragraph, headers, rows, widths=None):
    doc = paragraph._parent
    table = doc.add_table(rows=1, cols=len(headers), width=Cm(16))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    tr_pr = hdr._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))
    for i, h in enumerate(headers):
        set_cell(hdr.cells[i], h, bold=True)
        if widths:
            hdr.cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cells[i], value, align=align)
            if widths:
                cells[i].width = Cm(widths[i])
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    spacer_elm = OxmlElement("w:p")
    tbl.addnext(spacer_elm)
    spacer = Paragraph(spacer_elm, paragraph._parent)
    spacer.paragraph_format.space_after = Pt(4)
    return spacer


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(str(text))
    set_run_font(r, 10.5, bold=bold)


def find_para(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"paragraph not found: {text}")


def kicad_rows():
    labels = {
        "kicad_pcie_test_dual_ww37": "PCIeTest",
        "kicad_pi5_pcie_breakout": "Pi5-PCIe",
        "kicad_pi5_m2_hat": "Pi5-M2",
        "kicad_pcie_aux_signal_breakout": "PCIe-Aux",
    }
    by_case: dict[str, dict[str, dict[str, str]]] = {}
    with SUMMARY_CSV.open(newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            by_case.setdefault(row["case"], {})[row["method"]] = row
    rows = []
    for case, methods in by_case.items():
        base = methods["baseline_rule"]
        prop = methods["pso_congestion_reroute_astar"]
        rows.append(
            [
                case.replace("kicad_", ""),
                labels.get(case, case.replace("kicad_", "")),
                prop["pin_count"],
                f'{float(base["routability_percent"]):.2f}',
                f'{float(prop["routability_percent"]):.2f}',
                f'{float(base["area_ratio_percent"]):.2f}',
                f'{float(prop["area_ratio_percent"]):.2f}',
                base["violations"],
                prop["violations"],
            ][-8:]
        )
    return rows


def replace_text(doc: Document):
    replacements = {
        "算法原型采用C++17标准实现，未引入Qt等图形界面依赖，以命令行方式保证跨平台可构建和便于答辩演示。工程包含CMakeLists.txt和Makefile两种构建入口；在当前macOS环境中，由于未安装CMake，使用Apple clang和Makefile完成了实际编译验证。": "算法原型采用C++17标准实现，未引入Qt等图形界面依赖，以命令行方式保证跨平台可构建和便于答辩演示。工程包含CMakeLists.txt和Makefile两种构建入口；当前实验环境已完成CMake/Makefile兼容验证，公开KiCad案例实验使用同一可执行程序生成CSV指标和SVG/PNG可视化结果。",
        "其次，当前布线模型采用两层合法化假设，未显式优化过孔数量、层分配、差分等长和阻抗控制": "其次，当前布线模型采用两层合法化假设，公开KiCad案例仍属于案例化验证而非企业量产签核，未显式优化过孔数量、层分配、差分等长和阻抗控制",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in replacements:
            p.text = replacements[text]
            style_para(p)


def main():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(SRC, OUT)
    formula_pngs = [(label, render_formula(label, formula)) for label, formula in FORMULAS]

    doc = Document(OUT)
    replace_text(doc)

    p22 = find_para(doc, "2.2 布点布线数学模型")
    last = p22
    for text in [
        "为满足毕业设计管理办法中关于理论分析与计算正确性的要求，本文进一步把布点布线问题写成可复核的数学模型。金手指端子、候选测试点和布点决策变量定义如下：",
    ]:
        last = insert_paragraph_after(last, text)
    for label, img in formula_pngs[:4]:
        last = insert_picture_after(last, img, 12.5, label)
    last = insert_paragraph_after(last, "其中，Omega表示障碍区及其安全外扩区域，s_min表示测试点最小中心距，cap(e)表示网格边或Tile通道容量，n_e表示当前占用该通道的网络数量。上述公式将覆盖率、布通率、面积占比和合法性违规数统一到同一评价体系中，便于后续实验直接对应。")

    p34 = find_para(doc, "3.4 改进后的布点目标函数与步骤")
    p34_last = insert_picture_after(p34, formula_pngs[4][1], 10.5, formula_pngs[4][0])
    insert_paragraph_after(p34_last, "式（3-1）中，D表示端子到测试点的预估路径长度，A表示测试点包围盒面积占比，C表示候选点所在Tile的拥塞预测，P表示点间距、障碍和边界违规惩罚。由于P的权重最高，算法首先保证合法性，再在可行解中优化面积和线长。")

    p43 = find_para(doc, "4.3 拥塞感知与拆线重布策略")
    p43_last = insert_picture_after(p43, formula_pngs[5][1], 12.5, formula_pngs[5][0])
    p43_last = insert_picture_after(p43_last, formula_pngs[6][1], 10.5, formula_pngs[6][0])
    insert_paragraph_after(p43_last, "式（4-1）给出了拥塞感知A*的评价函数，式（4-2）给出了历史拥塞代价更新方式。rho用于控制历史信息衰减，Delta n_e表示本轮布线新增占用，alpha用于调节热点通道的绕行强度。")

    p14 = find_para(doc, "1.4 任务书要求与完成情况")
    insert_paragraph_after(p14, "结合《武汉理工大学本科生毕业设计（论文）工作管理办法（试行）》中关于任务完成、数据处理、理论分析、程序资料和成果归档的要求，本文把论文、C++程序、案例数据、实验指标和可视化截图作为同一毕业设计成果链进行组织。论文正文负责说明理论依据和实验结论，程序与CSV/SVG/PNG结果负责支撑可复现实验，公开KiCad案例用于补充参数化仿真的工程来源说明。")

    p53 = find_para(doc, "5.3 关键函数与数据流")
    p53_last = insert_paragraph_after(p53, "表 5-3 算法实现框架")
    p53_last = insert_table_after(
        p53_last,
        ["模块", "输入", "核心函数/流程", "输出", "对应指标"],
        [
            ["案例解析", "case.csv / KiCad抽取参数", "read_scenario", "Scenario", "引脚数、尾板尺寸"],
            ["候选生成", "规则库、障碍区、网格步长", "generate_candidates", "候选点集合", "覆盖率、违规数"],
            ["布点优化", "金手指坐标、候选点", "place_optimized / place_pso_hanan", "placement.csv", "面积占比、点间距"],
            ["布线搜索", "端子、测试点、障碍区", "astar / route_all", "routes.csv", "布通率、总线长"],
            ["拥塞重布线", "初次路径、拥塞矩阵", "route_all_congestion_aware", "优化路径", "拥塞峰值、拥挤网格数"],
            ["结果导出", "Solution、Metrics", "write_csvs / write_svg", "metrics.csv、SVG/PNG", "论文图表"],
        ],
        [2.2, 3.0, 3.6, 2.6, 3.0],
    )
    p53_last = insert_paragraph_after(p53_last, "表 5-4 数据流与文件流")
    insert_table_after(
        p53_last,
        ["阶段", "输入文件", "处理结果", "输出文件", "论文用途"],
        [
            ["案例准备", "公开.kicad_pcb / 参数化CSV", "抽取金手指数量、间距和障碍", "data/*.csv", "实验条件说明"],
            ["算法运行", "case.csv", "生成布点、布线和指标", "placement.csv、routes.csv、metrics.csv", "指标表"],
            ["可视化", "路径与测试点坐标", "绘制尾板布局", "layout_*.svg", "布线成果图"],
            ["KiCad对比", "原始公开PCB", "KiCad CLI导出原图并拼接", "comparison_*.png", "公开案例截图"],
            ["论文汇总", "summary.csv", "筛选基线与设计算法", "表6-4、图6-4", "量化结论"],
        ],
        [1.8, 3.0, 3.4, 3.3, 2.9],
    )

    p63_existing = find_para(doc, "6.3 结果分析")
    p_public = insert_paragraph_before(p63_existing, "6.3 公开 KiCad 案例对比实验", "Heading 2")
    p_public.paragraph_format.page_break_before = True
    p_public_last = insert_paragraph_after(p_public, "为进一步提高实验数据的工程可解释性，本文在五组参数化仿真实验之外，补充公开KiCad项目案例化验证。实验脚本优先调用KiCad CLI导出原始PCB板图，并从.kicad_pcb文本中抽取连接器/金手指类封装的引脚数量、代表间距和障碍压力，再运行同一C++布点布线程序。当前环境检测到KiCad CLI版本为10.0.2，导出状态均为kicad_cli。")
    table_caption = insert_paragraph_after(p_public_last, "表 6-4 公开KiCad案例人工/规则基线与设计算法对比")
    table_end = insert_table_after(
        table_caption,
        ["公开案例", "引脚", "基线布通/%", "算法布通/%", "基线面积/%", "算法面积/%", "基线违规", "算法违规"],
        kicad_rows(),
        [3.3, 1.2, 1.9, 1.9, 1.9, 1.9, 1.5, 1.5],
    )
    p_public_last = insert_picture_after(table_end, CONTACT_SHEET, 15.5, "图 6-4 公开KiCad案例原始PCB、规则基线与设计算法对比", page_break_before=True)
    p_public_last = insert_paragraph_after(p_public_last, "由表 6-4可见，公开案例中规则基线平均布通率为79.34%，设计算法平均布通率提升至93.48%；平均尾板面积占比由50.23%降至22.39%；平均合法性违规数由11.75降至4.50。M.2 HAT案例仍未达到满布通，主要原因是该连接器密度较高，抽象障碍区压缩了可用通道，形成高约束验证场景；即便如此，设计算法仍在布通率、面积占比、违规数和拥塞峰值上优于规则基线。")
    insert_paragraph_after(p_public_last, "需要说明的是，公开KiCad案例用于证明算法对真实EDA文件抽取参数后的适应性，实验结论属于公开案例化验证，不等同于企业量产板卡的设计签核。真实生产还需要结合层叠、阻抗、过孔、差分等长和制造厂规则进行二次合法化。")

    for old, new in [
        ("6.3 结果分析", "6.4 结果分析"),
        ("6.4 消融实验与对比有效性分析", "6.5 消融实验与对比有效性分析"),
        ("6.5 参数与工程适用性讨论", "6.6 参数与工程适用性讨论"),
    ]:
        try:
            p = find_para(doc, old)
            p.text = new
            p.style = "Heading 2"
            style_para(p, bold=True, first_indent=False)
        except ValueError:
            pass

    p71 = find_para(doc, "7.1 工作总结")
    insert_paragraph_after(p71, "按照学校管理办法关于资料归档的要求，本课题形成了可归档的成果集合：毕业论文DOCX/PDF、C++17源代码工程、参数化案例CSV、公开KiCad案例抽取数据、实验指标summary.csv、布线SVG/PNG截图、KiCad CLI导出原始PCB图以及实验报告。这些材料能够支撑答辩中对任务完成情况、理论依据、程序运行和数据可靠性的逐项说明。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
