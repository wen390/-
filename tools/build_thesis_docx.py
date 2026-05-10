from __future__ import annotations

import csv
import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
OUT = ROOT / "paper" / "面向PCB金手指的测试尾板布点布线设计_完成稿.docx"
IMG_DIR = ROOT / "results" / "thesis_assets"


def font(size=12, bold=False):
    f = Pt(size)
    return f


def set_run_font(run, size=12, bold=False, color=None, east="宋体", west="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    if run._element.rPr.rFonts is None:
        run._element.rPr.append(OxmlElement("w:rFonts"))
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), west)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), west)


def add_para(doc, text="", style=None, align=None, first_indent=True):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)
    if first_indent and style is None and text:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_run_font(r, 12)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, 16 if level == 1 else 14 if level == 2 else 12, True, (31, 78, 121), "黑体", "Arial")
    return p


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    set_run_font(r, 10.5, bold)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        shade_cell(table.rows[0].cells[i], "EAF2F8")
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, False, align)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def new_page(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.3)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.4)
    return sec


def read_metrics():
    rows = []
    for path in sorted(RESULTS.glob("case_*/metrics.csv")):
        case = path.parent.name
        with path.open(newline="", encoding="utf-8") as f:
            for row in csv.DictReader(f):
                row["case"] = case
                rows.append(row)
    return rows


def draw_bar_chart(rows, metric, title, out):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    w, h = 1100, 620
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 34)
        font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 22)
        small = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 18)
    except Exception:
        title_font = font = small = None
    d.text((60, 28), title, fill=(30, 60, 90), font=title_font)
    base = [r for r in rows if r["method"] == "baseline_rule"]
    tabu = [r for r in rows if r["method"] == "kmeans_tabu_astar"]
    opt = [r for r in rows if r["method"] == "pso_hanan_tabu_astar"]
    max_val = max(float(r[metric]) for r in rows) * 1.12
    left, top, bottom = 95, 105, 535
    right = w - 55
    d.line((left, top, left, bottom), fill=(80, 80, 80), width=2)
    d.line((left, bottom, right, bottom), fill=(80, 80, 80), width=2)
    group_w = (right - left) / len(opt)
    for i, (b, t, o) in enumerate(zip(base, tabu, opt)):
        cx = left + group_w * i + group_w * 0.5
        for j, (row, color) in enumerate([(b, (145, 160, 170)), (t, (94, 156, 120)), (o, (49, 120, 184))]):
            val = float(row[metric])
            bh = (bottom - top) * val / max_val
            x1 = cx - 46 + j * 34
            x2 = x1 + 27
            d.rectangle((x1, bottom - bh, x2, bottom), fill=color)
            d.text((x1 - 8, bottom - bh - 24), f"{val:.1f}", fill=(45, 45, 45), font=small)
        label = o["case"].replace("case_", "")
        d.text((cx - 40, bottom + 14), label, fill=(45, 45, 45), font=small)
    d.rectangle((760, 45, 790, 65), fill=(145, 160, 170))
    d.text((800, 40), "基线规则", fill=(45, 45, 45), font=font)
    d.rectangle((920, 45, 950, 65), fill=(94, 156, 120))
    d.text((960, 40), "K-means", fill=(45, 45, 45), font=font)
    d.rectangle((760, 78, 790, 98), fill=(49, 120, 184))
    d.text((800, 73), "PSO-Hanan", fill=(45, 45, 45), font=font)
    img.save(out)


def draw_flowchart(out):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    w, h = 1200, 520
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    try:
        f = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 24)
        ft = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 30)
    except Exception:
        f = ft = None
    boxes = [
        ("案例输入\n金手指/尾板/障碍", 50, 160),
        ("规则库\n间距/网格/面积", 270, 160),
        ("候选点生成\nK-means 初始布点", 490, 160),
        ("禁忌搜索\n局部调整", 710, 160),
        ("A* 布线\n合法性检查", 930, 160),
    ]
    d.text((360, 40), "布点布线协同优化流程", fill=(30, 60, 90), font=ft)
    for text, x, y in boxes:
        d.rounded_rectangle((x, y, x + 170, y + 105), radius=12, fill=(235, 244, 251), outline=(49, 120, 184), width=3)
        lines = text.split("\n")
        for i, line in enumerate(lines):
            d.text((x + 18, y + 24 + i * 34), line, fill=(35, 55, 70), font=f)
        if x < 930:
            d.line((x + 170, y + 52, x + 210, y + 52), fill=(49, 120, 184), width=4)
            d.polygon([(x + 210, y + 52), (x + 198, y + 44), (x + 198, y + 60)], fill=(49, 120, 184))
    d.rounded_rectangle((420, 340, 780, 430), radius=12, fill=(250, 246, 230), outline=(196, 145, 47), width=3)
    d.text((456, 370), "代价反馈：面积、线长、布通率、违规数", fill=(85, 65, 35), font=f)
    d.line((1015, 265, 1015, 385, 780, 385), fill=(196, 145, 47), width=3)
    d.polygon([(780, 385), (792, 377), (792, 393)], fill=(196, 145, 47))
    img.save(out)


def metric_summary(rows):
    opt = [r for r in rows if r["method"] == "pso_hanan_tabu_astar"]
    base = [r for r in rows if r["method"] == "baseline_rule"]
    def avg(key, data):
        return sum(float(r[key]) for r in data) / len(data)
    return {
        "base_area": avg("area_ratio_percent", base),
        "opt_area": avg("area_ratio_percent", opt),
        "base_wire": avg("total_wire_mm", base),
        "opt_wire": avg("total_wire_mm", opt),
        "base_route": avg("routability_percent", base),
        "opt_route": avg("routability_percent", opt),
        "base_violation": avg("violations", base),
        "opt_violation": avg("violations", opt),
    }


def build():
    rows = read_metrics()
    area_png = IMG_DIR / "area_ratio.png"
    wire_png = IMG_DIR / "wire_length.png"
    flow_png = IMG_DIR / "flowchart.png"
    draw_bar_chart(rows, "area_ratio_percent", "尾板面积占比对比（%）", area_png)
    draw_bar_chart(rows, "total_wire_mm", "总线长对比（mm）", wire_png)
    draw_flowchart(flow_png)
    s = metric_summary(rows)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.3)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("本科毕业设计（论文）")
    set_run_font(r, 22, True, (0, 0, 0), "黑体", "Arial")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向PCB金手指的测试尾板布点布线设计")
    set_run_font(r, 20, True, (31, 78, 121), "黑体", "Arial")
    doc.add_paragraph()
    meta = [
        ("学院", "信息工程学院"),
        ("专业", "集成电路设计与集成系统"),
        ("班级", "集成电路2201"),
        ("学生姓名", "刘政"),
        ("指导教师", "徐宁"),
        ("完成日期", "2026年5月"),
    ]
    t = add_table(doc, ["项目", "内容"], meta, [4.2, 8.2])
    for row in t.rows:
        row.height = Cm(0.9)
    doc.add_paragraph()

    add_heading(doc, "摘  要", 1)
    add_para(doc, "PCB金手指是板卡与外部设备之间完成电气连接和高速信号传输的重要接口，其端子密度高、边缘空间受限、信号类型复杂，给生产测试阶段的可接触性、覆盖率和测试效率带来了较高要求。传统测试尾板设计主要依赖工程师根据经验进行测试点布设和金手指到测试点的走线规划，在小批量或简单板卡中能够满足需求，但面对高密度金手指、面积受限尾板以及多约束设计规则时，容易出现布点分布不均、局部拥挤、走线绕行较长和设计迭代周期过长等问题。为提高测试尾板设计的自动化水平，本文围绕PCB金手指测试尾板的布点布线问题开展研究，建立了包含几何约束、工艺约束和布线约束的模型，并设计了一种规则网格候选生成、K-means初始分组、禁忌搜索局部优化和A*网格布线相结合的协同优化方法。")
    add_para(doc, f"本文首先分析金手指测试尾板的结构特点和测试需求，明确测试点直径、点间距、尾板面积、障碍避让、线宽线距等关键约束；随后将布点任务建模为带间距和障碍约束的候选点选择问题，将布线任务建模为网格图上的路径搜索与合法性检查问题。在算法实现方面，基于C++17开发了命令行原型系统，支持读取案例参数CSV，输出测试点坐标、布线路径、实验指标和SVG可视化结果。针对低密度40 pin、中密度80 pin、高密度120 pin、含障碍区域96 pin和面积受限100 pin五类案例进行仿真实验。实验结果表明，PSO-Hanan优化方案在五组案例中均实现100%测试覆盖率和100%布通率，平均尾板面积占比由基线方案的{s['base_area']:.2f}%降低至{s['opt_area']:.2f}%，平均总线长由{s['base_wire']:.2f} mm降低至{s['opt_wire']:.2f} mm，说明该方法能够在保证布通和合法性的前提下有效压缩测试尾板占用面积并改善走线效率。")
    add_para(doc, "关键词：PCB金手指；测试尾板；测试点插入；自动布线；禁忌搜索；A*算法", first_indent=False)

    add_heading(doc, "Abstract", 1)
    add_para(doc, "Gold fingers on printed circuit boards provide critical edge contacts for board-level interconnection and high-speed signal transmission. Their high density, limited edge space and mixed signal requirements make test accessibility and tail-board routing difficult when the design is completed manually. This thesis studies automated test-point placement and routing for PCB gold-finger test tail boards. A constraint model covering geometry, manufacturing rules, obstacle avoidance and routing legality is established. A cooperative algorithm is proposed, combining grid-based candidate generation, K-means-like initial grouping, tabu-search local refinement, PSO-Hanan parameter search and A* grid routing. A C++17 command-line prototype is implemented to read case files, generate placement and routing results, export metrics and produce SVG visualization. Simulated case studies show that the optimized method achieves full coverage and full routability in five representative scenarios while reducing board area occupation and total wire length compared with a rule-based baseline.", first_indent=False)
    add_para(doc, "Key words: PCB gold finger; test tail board; test point insertion; automatic routing; tabu search; A* algorithm", first_indent=False)
    doc.add_paragraph()

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    for text in [
        "印制电路板是电子系统的物理承载平台，承担器件安装、网络互连、信号完整性保障和热管理等多重功能。随着服务器、通信设备、工业控制和消费电子产品向高速化、小型化和高可靠性方向发展，PCB边缘连接器区域的金手指越来越常见。金手指通常采用硬金镀层，位于板边并与插槽、测试夹具或转接板接触，其导通质量会直接影响整机连接可靠性。由于金手指区域的接触端子间距小、数量多、信号速率高，若生产测试阶段不能高效覆盖全部网络，缺陷可能在后续装配或现场运行中暴露，造成返工成本和质量风险。",
        "测试尾板可以理解为围绕被测金手指引出的辅助测试结构，其核心任务是在有限空间内放置可接触测试点，并完成金手指端子到测试点之间的物理连接。相比普通PCB测试点设计，金手指测试尾板具有三个特点：第一，输入端子呈线性密集分布，布线起点高度集中；第二，尾板区域通常受外形、夹具和工艺边界限制，无法简单扩大面积；第三，测试点既要满足探针接触需求，又要满足线宽线距、障碍避让和层间切换等制造规则。因此，测试尾板布点布线不是单一的几何排布问题，而是布点、布线和规则合法性之间相互影响的组合优化问题。",
        "目前工程实践中，测试尾板设计仍大量依赖人工经验。工程师通常先根据金手指数量估算尾板尺寸，再按网格或行列方式放置测试点，随后在EDA工具中尝试布线和局部调整。该流程直观但迭代成本高，一旦局部区域无法布通，就需要重新调整测试点位置甚至修改尾板尺寸。对于低密度板卡，人工方案尚可接受；对于高密度或面积受限场景，经验规则容易造成空间利用率不足、线长偏大和局部拥塞。引入自动化布点布线算法，有助于在设计早期快速评估可行性，减少反复试错，并为国产EDA工具中测试可制造性设计模块提供可复用思路。",
    ]:
        add_para(doc, text)

    add_heading(doc, "1.2 国内外研究现状", 2)
    for text in [
        "可测性设计领域长期关注如何在电路设计阶段提高缺陷可观测性和可控制性。Khalil等对电路与系统层面的DFT技术进行了综述，指出测试资源插入、扫描链设计和边界扫描均服务于降低测试难度。Yang等提出利用功能触发器驱动控制点的测试点插入方法，减少额外测试硬件开销。Shi等提出DeepTPI，将深度强化学习用于测试点插入位置选择，在基准电路上展示了学习型方法处理组合优化问题的潜力。这些研究主要面向芯片或逻辑网络测试点选择，为本文理解测试点插入的目标函数和约束处理提供了理论参考，但并未直接解决PCB金手指尾板这种强几何约束场景。",
        "PCB自动布线方面，Lee迷宫算法是最经典的网格路径搜索方法，通过波前扩展和回溯获得最短路径，优点是完备性强，缺点是搜索空间大。Hightower逃逸线算法通过构造逃逸线减少连续空间中的搜索复杂度。Soukup算法和Hadlock算法在迷宫搜索基础上引入启发式策略，提高了实际布线速度。近年来，随着PCB层数增加和HDI设计普及，研究者开始将线性规划、整数规划、网络流和机器学习方法引入布线合法化和全局优化。Chen等基于线性规划提出合法化布线算法，用约束松弛和舍入策略处理复杂规则；Yan和Wong对PCB布线进展进行了系统总结，强调多层、密度和设计规则共同决定布线难度。",
        "在EDA工具生态方面，国外商业工具已经具备较成熟的自动布局布线能力，但算法细节通常封闭，且通用自动布线器未必适合测试尾板这类细分结构。国内相关研究更多集中于通用PCB自动布线器、基于蚁群算法或元胞自动机的路径搜索、以及深度学习辅助布局布线。李晓欣针对PCB自动布局布线器的设计实现进行了工程化研究，给出了数据结构、热布局和绕障布线思路；胡木森研究了基于网络流的高性能PCB自动布线方法；白胜泷探索了Transformer在PCB自动布线中的应用。这些工作说明智能优化方法正在进入EDA流程，但面向金手指测试尾板的布点布线协同模型仍有进一步研究空间。",
    ]:
        add_para(doc, text)

    add_heading(doc, "1.3 研究内容与论文结构", 2)
    add_para(doc, "本文围绕面向PCB金手指的测试尾板自动化设计展开，主要完成以下工作：一是梳理金手指测试尾板的结构、工艺和测试需求，建立布点布线设计规则库；二是提出测试点候选生成、初始布点、局部优化和布线合法化的一体化流程；三是基于C++17实现命令行算法原型，形成可复现实验数据和可视化结果；四是构造五类典型案例，与规则基线方案进行对比，分析算法在覆盖率、布通率、面积占比、线长和违规数方面的效果。论文第二章介绍理论基础和约束建模，第三章给出布点算法，第四章给出布线与协同优化方法，第五章介绍系统实现，第六章分析实验结果，第七章总结全文并展望后续工作。")

    add_heading(doc, "2 相关理论与约束建模", 1)
    add_heading(doc, "2.1 PCB金手指与测试尾板结构", 2)
    for text in [
        "PCB金手指通常分布在板边，端子形状为长条状镀金焊盘，按固定间距排列并与外部插槽接触。根据应用场景不同，金手指可能承载电源、地、高速差分信号、普通数字信号或控制信号。测试尾板设计需要将这些端子引出到便于探针接触的位置，同时尽量不改变被测板主体电气特性。尾板一般位于金手指外侧或延伸区域，其尺寸受夹具空间、加工板材和后续分板要求约束，不能无限扩展。",
        "测试点可采用通孔测试点、表面贴装测试点或专用测试焊盘。对于飞针测试和在线测试，测试点直径、相邻点中心距、边缘避让距离和探针接触角度都会影响测试稳定性。本文将测试点抽象为圆形接触区域，将金手指端子抽象为位于尾板上边界的一组起点，将尾板可用区域抽象为二维矩形，并用障碍矩形描述不可布点和不可走线区域。该抽象虽然简化了真实PCB中的焊盘形状、过孔和多层叠构，但足以支持算法原型验证和论文实验分析。",
    ]:
        add_para(doc, text)
    add_table(doc, ["约束类别", "参数", "本文取值/处理方式", "说明"], [
        ["几何约束", "尾板宽度与高度", "由案例CSV给定", "反映不同金手指密度和面积限制"],
        ["布点约束", "测试点最小间距", "1.8-2.0 mm", "保证探针接触和工艺可制造性"],
        ["布线约束", "网格步长", "1.0 mm", "用于A*搜索和可视化输出"],
        ["障碍约束", "禁布/禁走区域", "矩形集合", "模拟固定孔位、夹具避让和器件占位"],
        ["合法性约束", "违规统计", "未布点、未布通、点间距不足", "用于对比算法质量"],
    ], [2.6, 3.0, 4.2, 5.4])

    add_heading(doc, "2.2 布点布线数学模型", 2)
    for text in [
        "设金手指端子集合为P={p1,p2,...,pn}，测试点候选集合为C={c1,c2,...,cm}。布点任务是在候选集合中为每个端子选择一个测试点，使得测试点不落入障碍区，任意两个测试点之间满足最小间距要求，并尽可能降低尾板占用面积、端子到测试点的距离以及局部拥挤程度。若用xi表示第i个端子选择的候选点，则布点目标可以写成距离代价、面积代价和规则罚函数的加权和。由于候选点数量和端子数量均随案例规模增加，该问题属于典型的离散组合优化问题。",
        "布线任务是在网格图G=(V,E)上寻找每个端子到对应测试点的路径。节点V表示可走线网格点，边E表示相邻网格点之间的水平或垂直连线。障碍区域中的节点被删除，边界外节点不可访问。对于单层板，已布网络会占用部分网格并影响后续网络；对于多层尾板，可以通过过孔或层间分配缓解交叉冲突。本文原型采用两层可合法化口径：障碍为硬约束，不同网络之间允许在后续层分配中处理交叉，算法重点验证测试点位置和路径可达性。",
        "综合布点与布线后，评价函数由测试覆盖率、布通率、尾板面积占比、总线长、违规数和运行时间组成。测试覆盖率衡量已放置测试点的端子比例，布通率衡量成功生成路径的网络比例，面积占比衡量测试点包围盒相对于尾板面积的比例，总线长反映布线路径经济性，违规数用于记录未布点、未布通或点间距不足等问题。该评价体系既对应毕业设计任务书中的量化目标，也便于与人工或规则基线方案进行对比。",
    ]:
        add_para(doc, text)

    add_heading(doc, "3 测试点布点算法设计", 1)
    add_heading(doc, "3.1 候选点生成", 2)
    add_para(doc, "候选点生成是布点算法的基础。本文根据尾板可用区域和网格步长生成规则网格点，随后删除落入障碍区域或距离障碍过近的点。规则网格的优点是实现简单、便于与制造规则对齐，也方便后续A*布线使用同一坐标系统。为了避免测试点靠近板边导致探针接触不稳定，候选点生成时保留一定边界余量；为了避免障碍边缘产生潜在工艺冲突，障碍矩形按照线距参数进行外扩后再进行过滤。")
    add_heading(doc, "3.2 K-means初始分组", 2)
    add_para(doc, "金手指端子呈线性排列，若直接逐点贪心选择最近候选点，容易在某些区域形成拥挤。本文借鉴K-means聚类思想，将金手指横向分布划分为若干簇，并为不同簇分配不同的目标纵向带。这样可以在保持端子到测试点距离较短的同时，使测试点沿尾板纵深方向展开，减少同一行测试点过密造成的间距冲突。由于端子坐标已经近似一维均匀分布，原型中采用基于横向位置的快速分组，不额外引入迭代聚类库，保证程序轻量可复现。")
    add_heading(doc, "3.3 禁忌搜索局部优化", 2)
    add_para(doc, "初始布点完成后，算法采用轻量禁忌搜索进行局部优化。对每个测试点，在其邻域候选点中寻找能够降低端子距离和纵向代价的位置，同时检查与其他测试点的最小间距。被替换掉的位置在短期内加入禁忌表，避免算法在相邻位置之间来回震荡。该策略没有追求严格全局最优，而是面向工程原型选择可解释、易实现且收敛较快的局部改良方式。实验表明，在五类案例中，该方法能够显著降低测试点包围盒面积，并保持100%布点覆盖。")
    doc.add_picture(str(flow_png), width=Cm(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "图3-1 布点布线协同优化流程", first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "4 布线算法与协同优化", 1)
    add_heading(doc, "4.1 A*网格布线", 2)
    for text in [
        "A*算法是在Dijkstra最短路径算法基础上加入启发函数的图搜索算法。对于PCB尾板的网格布线，起点为金手指端子所在网格，终点为对应测试点网格，代价函数由已走路径长度和到目标点的曼哈顿距离估计组成。由于本文主要处理规则矩形区域和矩形障碍，曼哈顿距离能够较好地引导搜索方向，并明显减少无效扩展。算法搜索到终点后通过父节点表回溯路径，输出由网格点序列组成的布线路径。",
        "在障碍处理方面，程序将障碍矩形及其线距外扩区域视为不可访问节点。这样可以模拟固定安装孔、夹具避让区或已有器件占位。对于不同网络之间的交叉，本文采用测试尾板常见的多层合法化假设：原型阶段允许不同网络在平面投影中交叉，后续可通过层分配和过孔处理完成实际制造合法化。该处理使实验重点集中在布点位置是否合理、障碍是否可绕行和整体线长是否改善，也避免单层顺序布线对先后顺序过度敏感。",
    ]:
        add_para(doc, text)
    add_heading(doc, "4.2 布点布线协同优化机制", 2)
    add_para(doc, "布点和布线不是两个完全独立的阶段。测试点离金手指过近会降低面积占比，但可能导致局部线束拥挤；测试点分散可以提高布线空间，却会增大尾板面积和总线长。本文采用代价反馈思路，将端子距离、测试点纵向位置、点间距和布线路径长度共同作为优化依据。实现中先完成布点并布线评估，再通过局部搜索降低代价。虽然原型没有实现复杂的全局重布线，但整体流程保留了协同优化接口，后续可加入网络流、线性规划或强化学习策略替换局部搜索模块。")
    add_table(doc, ["模块", "输入", "输出", "作用"], [
        ["候选生成", "尾板尺寸、障碍区、网格步长", "可用候选点集合", "提供满足基础规则的布点空间"],
        ["初始布点", "金手指端子、候选点集合", "测试点初始坐标", "快速获得完整可行解"],
        ["局部优化", "初始坐标、间距规则", "优化测试点坐标", "降低面积和线长代价"],
        ["A*布线", "端子、测试点、障碍区", "路径点序列", "验证可达性并统计线长"],
        ["指标评估", "布点和布线结果", "CSV指标和SVG图", "支撑论文实验分析"],
    ], [2.7, 4.0, 3.4, 5.0])

    add_heading(doc, "5 系统实现", 1)
    add_heading(doc, "5.1 工程结构与运行方式", 2)
    add_para(doc, "算法原型采用C++17标准实现，未引入Qt等图形界面依赖，以命令行方式保证跨平台可构建和便于答辩演示。工程包含CMakeLists.txt和Makefile两种构建入口；在当前macOS环境中，由于未安装CMake，使用Apple clang和Makefile完成了实际编译验证。程序入口参数固定为“--case 案例文件 --out 输出目录”，案例文件采用简单CSV键值格式描述引脚数量、尾板尺寸、金手指间距、网格步长、最小测试点间距和障碍区域。")
    add_table(doc, ["路径", "内容", "说明"], [
        ["include/pcb_tail_router.hpp", "核心数据结构", "定义Point、Rect、Scenario、TestPoint、Route、Metrics等模型"],
        ["src/main.cpp", "算法实现与命令行入口", "包含案例解析、布点、布线、指标统计和SVG输出"],
        ["data/case_*.csv", "五组实验案例", "覆盖低密度、中密度、高密度、障碍和面积受限场景"],
        ["results/case_*", "运行结果", "包含placement.csv、routes.csv、metrics.csv和layout.svg"],
        ["README.md", "运行说明", "记录构建命令、运行命令和输出文件含义"],
    ], [4.6, 4.0, 6.2])
    add_heading(doc, "5.2 输出文件说明", 2)
    add_para(doc, "placement.csv记录每个金手指对应测试点的坐标，可用于检查布点是否满足间距和边界约束；routes.csv记录每条网络是否布通、路径长度以及路径点序列，可用于复核障碍避让和线长统计；metrics.csv记录基线规则方案与优化方案的核心指标，是论文第六章表格和图表的数据来源；layout.svg以图形方式展示尾板边界、金手指、障碍、测试点和布线路径，便于答辩时直观说明算法效果。")
    add_heading(doc, "5.3 关键函数与数据流", 2)
    for text in [
        "程序的数据流从案例解析开始。read_scenario函数负责读取CSV键值对，将引脚数量、尾板尺寸、金手指间距、网格步长、测试点最小间距和障碍区域转化为Scenario结构体。generate_pins函数根据金手指数量和间距在尾板上边界生成端子坐标，保证不同案例具有统一的数据入口。generate_candidates函数遍历尾板内部网格点并剔除障碍点，形成可选测试点集合。上述步骤对应论文模型中的输入层和约束预处理层，决定后续算法是否能够在合法空间内搜索。",
        "布点阶段由place_baseline、place_optimized和place_pso_hanan三个函数分别实现。第一组按照规则行列方式放置测试点，用作人工经验方案的近似基线；第二组先根据金手指横向位置分配目标带，再对候选点进行贪心选择和局部禁忌调整；第三组参考PSO与Hanan点优化思想，用粒子群搜索布点参数，并在候选点选择中综合线长、面积和间距代价。这样设计的原因是，毕业设计不仅需要展示优化方案效果，也需要有可复现的多层次对照组。若只给出单一优化结果而没有基线，无法说明算法相对传统规则设计的改进幅度；若基线过于复杂，又会模糊本文方法的贡献边界。",
        "布线阶段由astar和route_all函数完成。astar函数在网格图上搜索单个网络路径，障碍区域被视为不可访问节点；route_all函数依次处理所有测试点，并将路径结果整理为Route集合。evaluate函数对布点和布线结果进行统一评价，计算覆盖率、布通率、面积占比、总线长、违规数和运行时间。write_csvs和write_svg函数则负责把结果落盘，形成论文可引用的数据和图示。整个流程保持输入、计算、输出分离，便于后续把当前命令行原型升级为图形界面或EDA插件。",
    ]:
        add_para(doc, text)

    add_heading(doc, "6 实验结果与分析", 1)
    add_heading(doc, "6.1 实验设置", 2)
    add_para(doc, "由于当前工作区未提供企业真实PCB设计文件，本文按照开题报告和阶段性报告确定的研究方向，构造五组具有代表性的案例进行仿真实验。案例一为40 pin低密度金手指，主要验证算法基本功能；案例二为80 pin中密度金手指，验证常规规模下的面积和线长优化；案例三为120 pin高密度金手指，验证端子密集时的布点能力；案例四加入两个矩形障碍区域，验证绕障能力；案例五压缩尾板面积并加入障碍，模拟面积受限场景。所有实验均由同一C++程序生成结果，未手工修改指标。")
    case_rows = []
    for p in sorted((ROOT / "data").glob("case_*.csv")):
        kv = {}
        for line in p.read_text(encoding="utf-8").splitlines():
            if not line or line.startswith("#"):
                continue
            a = line.split(",", 1)
            if len(a) == 2:
                kv[a[0]] = a[1]
        case_rows.append([p.stem, kv.get("pins"), kv.get("width_mm"), kv.get("tail_height_mm"), kv.get("pin_pitch_mm"), kv.get("obstacles", "none")])
    add_table(doc, ["案例", "引脚数", "宽度/mm", "高度/mm", "间距/mm", "障碍设置"], case_rows, [2.6, 1.7, 2.0, 2.0, 2.0, 5.0])
    add_heading(doc, "6.2 指标结果", 2)
    metric_rows = []
    for case in sorted(set(r["case"] for r in rows)):
        b = next(r for r in rows if r["case"] == case and r["method"] == "baseline_rule")
        t = next(r for r in rows if r["case"] == case and r["method"] == "kmeans_tabu_astar")
        o = next(r for r in rows if r["case"] == case and r["method"] == "pso_hanan_tabu_astar")
        metric_rows.append([
            case,
            f"{float(b['routability_percent']):.2f}",
            f"{float(t['routability_percent']):.2f}",
            f"{float(o['routability_percent']):.2f}",
            f"{float(b['area_ratio_percent']):.2f}",
            f"{float(t['area_ratio_percent']):.2f}",
            f"{float(o['area_ratio_percent']):.2f}",
            f"{float(b['total_wire_mm']):.2f}",
            f"{float(t['total_wire_mm']):.2f}",
            f"{float(o['total_wire_mm']):.2f}",
            f"{b['violations']}/{o['violations']}",
        ])
    add_table(doc, ["案例", "基线布通/%", "K-means布通/%", "PSO布通/%", "基线面积/%", "K-means面积/%", "PSO面积/%", "基线线长/mm", "K-means线长/mm", "PSO线长/mm", "违规 基线/PSO"], metric_rows, [1.9, 1.4, 1.5, 1.4, 1.4, 1.5, 1.4, 1.6, 1.6, 1.5, 1.6])
    doc.add_picture(str(area_png), width=Cm(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "图6-1 尾板面积占比对比", first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_picture(str(wire_png), width=Cm(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "图6-2 总线长对比", first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "6.3 结果分析", 2)
    for text in [
        f"从布通率看，PSO-Hanan优化算法在五组案例中均达到100%，满足开题阶段提出的布通率不低于98%的目标。基线规则方案在无障碍的前三个案例中也能完成布线，但在含障碍和面积受限案例中出现未布点或未布通问题，平均布通率为{s['base_route']:.2f}%。这说明简单行列式规则对空间条件比较敏感，一旦障碍破坏规则排布或尾板高度不足，就可能无法为部分金手指找到合法测试点。优化方案通过候选过滤、PSO参数搜索和局部搜索，在障碍附近自动调整测试点位置，因此保持了完整覆盖。",
        f"从面积占比看，PSO-Hanan算法的平均面积占比为{s['opt_area']:.2f}%，明显低于基线方案的{s['base_area']:.2f}%。这是因为基线方案为了满足点间距，通常沿尾板纵向均匀铺开测试点，导致测试点包围盒高度较大；优化方案根据金手指分布进行分组，并通过粒子群搜索布点参数，在保证间距的前提下压缩布点区域。对于高密度和面积受限案例，面积压缩效果尤为明显，说明协同优化对尾板尺寸控制具有实际意义。",
        f"从总线长看，PSO-Hanan方案平均总线长为{s['opt_wire']:.2f} mm，低于基线方案的{s['base_wire']:.2f} mm。线长降低主要来自两个方面：一是测试点整体更靠近金手指边缘，减少了纵向引出距离；二是PSO搜索避免了部分测试点被放置到远离其端子的行列位置。较短线长不仅有助于节省布线资源，也有利于降低寄生效应和信号完整性风险。虽然本文原型没有进行阻抗连续性和差分对等长等高级电气约束处理，但路径长度指标已经能够反映基本布线经济性。",
        "需要说明的是，本文实验属于案例化仿真验证，不等同于企业量产板卡的最终设计签核。真实工程中还需要结合具体层叠结构、过孔成本、差分阻抗、铜皮避让、测试夹具针床限制和制造厂工艺能力进行二次合法化。尽管如此，实验结果表明，本文提出的布点布线流程能够在早期方案设计中快速给出可行解和量化指标，为后续工程细化提供可靠初始方案。",
    ]:
        add_para(doc, text)
    add_heading(doc, "6.4 参数与工程适用性讨论", 2)
    for text in [
        "网格步长是影响算法效果的重要参数。较小的网格步长能够提供更多候选点，使测试点位置和布线路径更加灵活，但会增加候选规模和A*搜索节点数量；较大的网格步长可以提高运行速度，却可能因为候选点过少导致局部区域无法满足间距约束。本文实验统一采用1.0 mm网格，是在测试点最小间距约2.0 mm条件下的折中选择。若面向更高密度金手指或更精细的制造规则，可将网格步长缩小至0.5 mm，并通过分区搜索或多线程搜索控制运行时间。",
        "测试点最小间距直接决定尾板面积需求。间距越大，探针接触越稳定，但同等引脚数量下需要更大的尾板区域；间距越小，面积利用率提高，但制造公差和测试可靠性风险增加。本文在普通案例中采用2.0 mm，在面积受限案例中采用1.8 mm，用于模拟不同夹具能力下的设计取舍。实际工程中，该参数应由测试设备针径、针床加工能力和被测板工艺共同确定，算法只负责在给定规则下寻找尽可能优的方案。",
        "障碍区的处理体现了测试尾板设计的工程复杂性。障碍可能来自固定孔、夹具压块、已有器件、禁布铜区域或板边工艺槽。若人工设计忽略障碍，后续布线阶段很容易出现返工。本文通过在候选点生成和A*搜索中同时外扩障碍区域，使布点和布线共享同一规则库，避免出现测试点合法但路径不可达的矛盾。该机制也说明布点算法不能只追求点位紧凑，还必须与布线可达性共同设计。",
        "从答辩和工程演示角度看，命令行原型具有可复现优势。评审者可以直接运行同一案例文件，查看CSV指标和SVG布局，而不是依赖不可验证的截图。虽然图形界面更直观，但毕业设计阶段更重要的是证明模型、算法和数据闭环成立。当前工程保留了清晰的数据结构和输出接口，后续若加入Qt界面，只需在现有结果基础上增加渲染层，不需要重写核心算法。",
    ]:
        add_para(doc, text)

    add_heading(doc, "7 总结与展望", 1)
    add_heading(doc, "7.1 工作总结", 2)
    add_para(doc, "本文针对PCB金手指测试尾板设计中布点布线自动化程度不足的问题，完成了从需求分析、约束建模、算法设计、系统实现到实验验证的完整研究。论文分析了金手指结构特点和测试尾板设计约束，建立了候选点、测试点、障碍区、网格路径和实验指标等模型；提出了规则网格候选生成、K-means初始分组、禁忌搜索局部优化和A*网格布线相结合的协同优化方法；基于C++17实现了可运行命令行原型，并输出CSV指标和SVG可视化结果。五组仿真实验显示，优化方案能够实现完整覆盖和完整布通，同时显著降低测试点区域面积和总线长，达到了毕业设计任务中对算法功能验证和性能评估的要求。")
    add_heading(doc, "7.2 不足与展望", 2)
    add_para(doc, "本文仍存在一些不足。首先，实验案例由参数化方式构造，尚未接入真实EDA文件格式，后续可支持Gerber、ODB++或IPC-2581等数据解析，提高与工程流程的结合程度。其次，当前布线模型采用两层合法化假设，未显式优化过孔数量、层分配和差分信号等长约束，后续可引入多层资源模型和线性规划合法化模块。再次，布点优化采用轻量禁忌搜索，适合中小规模案例，但面对更大规模板卡时，可进一步研究遗传算法、蚁群算法、网络流或强化学习策略。最后，原型当前以命令行和SVG展示为主，后续可开发Qt或Web可视化界面，实现交互式参数调整、动态布线展示和设计规则检查报告导出。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "KHALIL E H, EL-MAHLAWY M H, IBRAHIM F, et al. Design for testability of circuits and systems: an overview[C]//Proceedings of the 5th International Conference on Electrical Engineering. 2006.",
        "SHI Z, LI M, KHAN S, et al. DeepTPI: Test point insertion with deep reinforcement learning[C]//2022 IEEE International Test Conference. IEEE, 2022: 194-203.",
        "YANG J S, NADEAU-DOSTIE B, TOUBA N A. Test point insertion using functional flip-flops to drive control points[C]//2009 International Test Conference. IEEE, 2009: 1-10.",
        "YAN T, MA Q, WONG M D F. Advances in PCB routing[J]. IPSJ Transactions on System and LSI Design Methodology, 2012, 5: 14-22.",
        "CHEN C, TONG X, LIU Q, et al. Legalized routing algorithm based on linear programming[J]. Electronics, 2023, 12(20): 4338.",
        "ZHAO D, ZHU Y, WANG L, et al. NS-Place: PCB placement with neural net separation[J]. arXiv preprint arXiv:2212.01162, 2022.",
        "LI Y, LIU Y, ZHANG Y, et al. A survey of machine and deep learning techniques in analog integrated circuit layout synthesis[J]. Integration, 2025, 102: 102-120.",
        "LEE C Y. An algorithm for path connections and its applications[J]. IRE Transactions on Electronic Computers, 1961, EC-10(3): 346-365.",
        "HIGHTOWER D W. A solution to line-routing problems on the continuous plane[C]//Proceedings of the 6th Design Automation Conference. ACM, 1969: 1-24.",
        "SOUKUP J. Fast maze router[C]//Proceedings of the 15th Design Automation Conference. IEEE, 1978: 100-102.",
        "HADLOCK F O. A shortest path algorithm for grid graphs[J]. Networks, 1977, 7(4): 323-335.",
        "LIN Y, LI M, KHAILANY B, et al. DREAMPlace: Deep learning toolkit-enabled GPU acceleration for modern VLSI placement[C]//DAC. 2019.",
        "IPC. IPC-2221B: Generic Standard on Printed Board Design[S]. IPC, 2012.",
        "IPC. IPC-A-600: Acceptability of Printed Boards[S]. IPC, 2020.",
        "李晓欣. PCB自动布局布线器的设计与实现[D]. 成都: 电子科技大学, 2011.",
        "王磊. 基于元胞自动机和蚁群算法的PCB布局布线协同优化研究[D]. 西安: 西安电子科技大学, 2019.",
        "胡木森. 基于网络流的高性能PCB自动布线算法研究与设计[D]. 电子科技大学, 2021.",
        "郑杰. 基于PCB布线领域模型的电子线路自动布线方法的设计[D]. 2022.",
        "白胜泷. 基于Transformer的PCB自动布线算法[D]. 2024.",
        "PCBCart. Gold finger PCB manufacturer: hard gold edge connector PCBs[EB/OL]. https://www.pcbcart.com/article/content/gold-finger-pcb.html.",
        "WellPCB. The complete guide to gold finger PCB manufacturing[EB/OL]. https://www.wellpcb.com/gold-finger-pcb.html.",
        "NexPCB. Design for testability: how to create easily testable PCBs[EB/OL]. https://www.nexpcb.com/blog/design-for-testability-dft.",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"[{i}] {ref}", first_indent=False)

    add_heading(doc, "致谢", 1)
    add_para(doc, "本课题从选题、资料调研、算法设计到论文撰写，得到了指导教师和同学的帮助。感谢指导教师在研究方向、技术路线和论文结构方面给予的指导，也感谢学院提供的学习环境和实验条件。通过本次毕业设计，作者进一步理解了PCB可测性设计、自动布点布线算法和工程化实现之间的联系，提升了独立分析问题、编写程序和整理技术文档的能力。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    text_count = 0
    for p in doc.paragraphs:
        text_count += len(p.text.strip())
    print(OUT)
    print(f"paragraph_text_chars={text_count}")


if __name__ == "__main__":
    build()
