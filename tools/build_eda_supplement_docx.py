from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
RESULT_DIR = ROOT / "results" / "eda_cases"
OUT = ROOT / "面向PCB金手指的测试尾板布点布线设计_EDA解析补充修改稿.docx"


def set_run_font(run, size=12, bold=False, color=None, east="宋体", west="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:ascii"), west)
    rfonts.set(qn("w:hAnsi"), west)


def style_para(p, first_indent=True, align=None, size=12):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    for r in p.runs:
        set_run_font(r, size=size)


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=18 if level == 1 else 15, bold=True, east="黑体", west="Times New Roman")
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 12)
    style_para(p)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, 12)
    return p


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(15)
    r = p.add_run(str(text))
    set_run_font(r, 9.5, bold=bold)


def shade_cell(cell, fill="EAF2F8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i])
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cells[i], value, align=align)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_picture(doc, path, width_cm, caption):
    add_caption(doc, caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = 1
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    return p


def add_case_image_grid(doc, case_name, case_dir, figure_no):
    add_caption(doc, f"图A-{figure_no} {case_name}案例独立截图对比")
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    items = [
        ("原始PCB", case_dir / "01_original_pcb.png"),
        ("解析覆盖", case_dir / "02_parse_overlay.png"),
        ("规则基线", case_dir / "02_baseline_rule.png"),
        ("设计算法", case_dir / "03_algorithm_pso_congestion_reroute_astar.png"),
    ]
    for idx, (label, path) in enumerate(items):
        cell = table.rows[idx // 2].cells[idx % 2]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.width = Cm(8.1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        r = p.add_run(label)
        set_run_font(r, 9.5, bold=True, east="黑体", west="Times New Roman")
        if path.exists():
            pic = cell.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.paragraph_format.first_line_indent = None
            pic.add_run().add_picture(str(path), width=Cm(7.8))
    doc.add_paragraph()


def read_summary():
    rows = list(csv.DictReader((RESULT_DIR / "summary.csv").open(encoding="utf-8")))
    by_case: dict[str, dict[str, dict]] = {}
    for row in rows:
        by_case.setdefault(row["case"], {})[row["method"]] = row
    return by_case


def source_rows(by_case):
    rows = []
    for case, methods in by_case.items():
        one = next(iter(methods.values()))
        rows.append([
            case.replace("eda_", ""),
            one["repo"],
            one["license"],
            Path(one["eda_file"]).name,
            one["footprint"],
            one["pins"],
            f'{one["gerber_files"]}/{one["drill_files"]}',
            one["drill_hits"],
        ])
    return rows


def metric_rows(by_case):
    rows = []
    for case, methods in by_case.items():
        base = methods["baseline_rule"]
        final = methods["pso_congestion_reroute_astar"]
        rows.append([
            case.replace("eda_", ""),
            base["routability_percent"],
            final["routability_percent"],
            base["area_ratio_percent"],
            final["area_ratio_percent"],
            base["violations"],
            final["violations"],
        ])
    return rows


def case_display_name(case):
    return case.replace("eda_", "").replace("_", " ")


def separate_case_dirs(by_case):
    base = RESULT_DIR / "separate_screenshots"
    rows = []
    for case in by_case:
        suffix = case.replace("eda_", "")
        matches = sorted(base.glob(f"eda_case_*_{suffix}"))
        if matches:
            rows.append((case_display_name(case), matches[0]))
    return rows


def make_flowchart() -> Path:
    path = RESULT_DIR / "eda_parse_flow.png"
    img = Image.new("RGB", (1280, 360), "white")
    draw = ImageDraw.Draw(img)
    boxes = [
        ("公开EDA文件\n.kicad_pcb / Gerber / DRL", 40),
        ("格式识别\nKiCad S-expression\nGerber/XNC坐标", 270),
        ("结构提取\n板框/封装/焊盘/钻孔", 500),
        ("约束映射\n端子坐标/障碍/规则", 730),
        ("算法运行\nCSV -> C++ -> SVG/PNG", 960),
    ]
    for text, x in boxes:
        draw.rounded_rectangle((x, 95, x + 200, 235), radius=12, fill="#eef4fb", outline="#446b91", width=2)
        for i, line in enumerate(text.split("\n")):
            draw.text((x + 18, 120 + i * 28), line, fill="#1f2d3d")
        if x < 960:
            draw.line((x + 205, 165, x + 250, 165), fill="#446b91", width=3)
            draw.polygon([(x + 250, 165), (x + 238, 157), (x + 238, 173)], fill="#446b91")
    draw.text((40, 300), "输出：data/eda_cases/*.csv、results/eda_cases/summary.csv、解析覆盖图、三联对比截图、论文补充材料", fill="#333")
    img.save(path)
    return path


def main():
    by_case = read_summary()
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    add_heading(doc, "真实EDA文件解析补充修改稿", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run("用于《面向PCB金手指的测试尾板布点布线设计》审阅意见回应")
    set_run_font(r, 12)

    add_heading(doc, "1 修改说明", 2)
    add_para(doc, "针对审阅意见中关于真实EDA文件解析的建议，补充实现了公开KiCad原生板文件、Gerber层文件和Excellon钻孔文件的轻量级解析流程。解析结果不再只停留在未来工作中，而是形成可运行的数据链：公开EDA文件解析为算法CSV输入，C++布点布线程序输出指标与可视化截图，论文补充稿给出解析依据、实验案例、量化结果和截图证据。")
    add_para(doc, "本补充稿可作为第5章系统实现、第6章实验结果和第7章总结展望的局部替换材料。为避免夸大结论，文中统一使用“公开EDA案例化验证”表述，不声称来自企业量产数据。")

    add_heading(doc, "2 建议加入第5章：真实EDA文件解析流程", 2)
    add_picture(doc, make_flowchart(), 15.8, "图A-1 真实EDA文件解析与算法实验数据流")
    add_para(doc, "新增脚本 tools/eda_case_experiments.py 以KiCad原生板文件为主要输入，按照KiCad官方S-expression板文件结构提取footprint、pad、Edge.Cuts板框和连接器候选；当公开项目同时提供Gerber/Excellon制造文件时，脚本进一步解析Gerber层文件数量、Edge_Cuts边界和钻孔命中数，用于证明制造文件级信息已经进入实验报告。")
    add_table(doc, ["解析对象", "提取内容", "输出字段", "用途"], [
        [".kicad_pcb", "板框、footprint、pad坐标、连接器候选", "pin_positions、footprint、pins", "生成真实端子坐标输入"],
        ["Gerber", "层文件、Edge_Cuts/Profile边界、Gerber元数据", "gerber_files、gerber_bbox", "补充制造文件解析证据"],
        ["Excellon/DRL", "钻孔文件、钻孔坐标命中数", "drill_files、drill_hits", "说明通孔/安装孔等制造信息"],
        ["C++ CSV", "端子坐标、尾板尺寸、障碍区、规则参数", "data/eda_cases/*.csv", "驱动布点布线算法"],
    ], [3.0, 5.2, 3.8, 4.0])

    add_heading(doc, "3 建议加入第6章：公开EDA案例化验证", 2)
    add_para(doc, "本次补充实验共使用7个公开EDA案例，其中4个为原有公开KiCad案例，3个为新增GitHub公开项目，覆盖PCIe、mini-PCIe、M.2/NGFF等边缘连接器或金手指类结构。实验环境检测到KiCad CLI版本为10.0.2，所有案例均完成原始板图导出、解析覆盖图生成和算法布点布线截图。")
    add_table(doc, ["案例", "公开仓库", "许可证", "EDA文件", "选中封装", "引脚", "Gerber/DRL", "钻孔数"], source_rows(by_case), [2.3, 2.8, 1.5, 2.3, 2.9, 0.9, 1.3, 0.9])
    add_table(doc, ["案例", "基线布通/%", "算法布通/%", "基线面积/%", "算法面积/%", "基线违规", "算法违规"], metric_rows(by_case), [3.0, 2.0, 2.0, 2.0, 2.0, 1.6, 1.6])
    add_para(doc, "为避免总览图在论文版面中缩放过小，以下将7个公开EDA案例拆分为独立截图组。每组均包含原始PCB板图、解析覆盖图、规则基线结果和设计算法结果，可作为论文实验截图或答辩材料直接引用。")
    for idx, (case_name, case_dir) in enumerate(separate_case_dirs(by_case), start=2):
        doc.add_page_break()
        add_case_image_grid(doc, case_name, case_dir, idx)
    add_para(doc, "由补充实验可见，规则基线在高密度M.2/NGFF案例中出现明显未布通和违规，设计算法在7个公开EDA案例中均达到100%布通；平均布通率由72.94%提升至100.00%，平均尾板面积占比由72.30%降低至27.80%。该结果说明，新增真实EDA解析入口能够把公开EDA文件中的端子结构转化为算法可处理的输入，并形成可复现实验闭环。")

    add_heading(doc, "4 建议替换第7.2展望相关表述", 2)
    add_para(doc, "原文“尚未接入真实EDA文件格式”建议替换为：本文已完成KiCad原生板文件和Gerber/Excellon制造文件的初步解析，实现了从公开EDA文件到算法CSV输入、布点布线结果和截图证据的闭环。受本科毕业设计周期和企业数据获取条件限制，当前解析仍以公开KiCad工程和轻量制造文件字段为主，尚未覆盖ODB++、IPC-2581等工业交换格式，也未完整处理差分阻抗、层叠规则和制造厂全套DRC约束。后续可在现有解析接口基础上继续扩展ODB++/IPC-2581导入、多层合法化、过孔成本、差分等长约束和GUI交互展示。")

    add_heading(doc, "5 可补充参考资料", 2)
    refs = [
        "KiCad Developer Documentation. Board File Format: https://dev-docs.kicad.org/en/file-formats/sexpr-pcb/",
        "KiCad Documentation. Command-Line Interface: https://docs.kicad.org/master/en/cli/cli.html",
        "Ucamco. Gerber Format Specifications: https://www.ucamco.com/en/guest/downloads/gerber-format",
        "mengstr/PCIeTestPCB: https://github.com/mengstr/PCIeTestPCB",
        "m1geo/Pi5_PCIe: https://github.com/m1geo/Pi5_PCIe",
        "mithro/kicad-mini-pci-express: https://github.com/mithro/kicad-mini-pci-express",
        "serg987/coral-dual-m2-adapter-pcb: https://github.com/serg987/coral-dual-m2-adapter-pcb",
        "themainframe/5g-m2-usb3-interface-pcb: https://github.com/themainframe/5g-m2-usb3-interface-pcb",
    ]
    for ref in refs:
        add_para(doc, ref)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
